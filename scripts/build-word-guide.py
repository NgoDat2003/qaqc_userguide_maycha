from __future__ import annotations

import json
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
CONTENT_JSON = ROOT / "src" / "content" / "guide-content.json"
PUBLIC_DIR = ROOT / "public"
DOWNLOAD_DIR = PUBLIC_DIR / "downloads"
OUTPUT_NAME = "huong-dan-qaqc.docx"
DISPLAY_REPLACEMENTS = (
    (re.compile(r"\bUAT\b", re.IGNORECASE), "hệ thống"),
    (re.compile(r"\bledger\b", re.IGNORECASE), "danh sách hình ảnh"),
    (re.compile(r"\bprivate\b", re.IGNORECASE), "nội bộ"),
)


def sanitize_display(value):
    if isinstance(value, str):
        for pattern, replacement in DISPLAY_REPLACEMENTS:
            value = pattern.sub(replacement, value)
        return value
    if isinstance(value, list):
        return [sanitize_display(item) for item in value]
    if isinstance(value, dict):
        return {key: sanitize_display(item) for key, item in value.items()}
    return value


def slugify(text: str, fallback: str) -> str:
    normalized = unicodedata.normalize("NFD", text.lower())
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "-", ascii_text).strip("-")
    return slug[:48].strip("-") or fallback


def bookmark_name(text: str, fallback: str) -> str:
    return "toc_" + slugify(text, fallback).replace("-", "_")


def set_run_font(run, size: int | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str, *, bold: bool = False):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    props.append(color)
    if bold:
        props.append(OxmlElement("w:b"))

    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text

    run.append(props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_pageref_field(paragraph, anchor: str, fallback_page: int):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f" PAGEREF {anchor} \\h ")

    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = str(fallback_page)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_text_paragraph(doc: Document, text: str, *, style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, 11, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_list(doc: Document, items: list[str], style: str):
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        run = table.rows[0].cells[index].paragraphs[0].add_run(header)
        set_run_font(run, 10, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(headers)]):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, 10)
    doc.add_paragraph()


def add_image(doc: Document, block: dict):
    image_path = PUBLIC_DIR / block["src"].lstrip("/")
    if not image_path.exists():
        return

    try:
        with Image.open(image_path) as image:
            width, height = image.size
    except Exception:
        width, height = 1200, 800

    max_width = 6.4
    if height > width:
        max_width = 3.2
    elif width < 900:
        max_width = 5.0

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(max_width))

    caption = block.get("caption") or block.get("alt")
    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, 9, italic=True)
        caption_run.font.color.rgb = RGBColor(85, 96, 115)


def estimate_toc_pages(sections: list[dict]) -> dict[str, int]:
    page = 3
    units_on_page = 0.0
    result: dict[str, int] = {}

    for section in sections:
        result[section["id"]] = page
        units = 0.18 if section["level"] == 1 else 0.12
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                units += max(0.08, len(block["text"]) / 900)
            elif block["type"] == "label":
                units += 0.08
            elif block["type"] in {"steps", "bullets"}:
                units += 0.08 * len(block["items"])
            elif block["type"] == "table":
                units += 0.14 + 0.06 * len(block["rows"])
            elif block["type"] == "image":
                image_path = PUBLIC_DIR / block["src"].lstrip("/")
                try:
                    with Image.open(image_path) as image:
                        width, height = image.size
                    units += 0.78 if height > width else 0.56
                except Exception:
                    units += 0.55
        units_on_page += units
        while units_on_page >= 1:
            page += 1
            units_on_page -= 1
    return result


def add_cover(doc: Document, guide: dict):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(guide["title"])
    set_run_font(title_run, 22, bold=True)
    title_run.font.color.rgb = RGBColor(12, 23, 43)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version.add_run(guide["version"])
    set_run_font(version_run, 13, bold=True)
    version_run.font.color.rgb = RGBColor(199, 160, 59)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(guide["scope"])
    set_run_font(subtitle_run, 11)

    add_text_paragraph(doc, f"Đối tượng sử dụng: {guide['audience']}", bold=True)
    add_text_paragraph(doc, guide["notice"])
    add_text_paragraph(doc, f"Tổng số mục: {len(guide['sections'])}")
    add_text_paragraph(doc, f"Tổng số hình ảnh: {guide['imageCount']}")
    doc.add_page_break()


def add_manual_toc(doc: Document, sections: list[dict], anchors: dict[str, str]):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("MỤC LỤC")
    set_run_font(heading_run, 16, bold=True)

    estimated_pages = estimate_toc_pages(sections)
    for section in sections:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.28 if section["level"] == 2 else 0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.25),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        add_internal_link(
            paragraph,
            section["title"],
            anchors[section["id"]],
            bold=section["level"] == 1,
        )
        tab = paragraph.add_run("\t")
        set_run_font(tab, 11)
        add_pageref_field(paragraph, anchors[section["id"]], estimated_pages[section["id"]])
    doc.add_page_break()


def mark_fields_dirty(doc: Document):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_document():
    guide = sanitize_display(json.loads(CONTENT_JSON.read_text(encoding="utf-8")))
    anchors = {
        section["id"]: bookmark_name(section["id"], f"section_{index + 1}")
        for index, section in enumerate(guide["sections"])
    }

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in (("Title", 22, "0C172B"), ("Heading 1", 16, "0C172B"), ("Heading 2", 13, "274C77")):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 2"].paragraph_format.keep_with_next = True

    add_cover(doc, guide)
    add_manual_toc(doc, guide["sections"], anchors)

    for index, section_data in enumerate(guide["sections"], start=1):
        heading_style = "Heading 1" if section_data["level"] == 1 else "Heading 2"
        heading = doc.add_paragraph(style=heading_style)
        run = heading.add_run(section_data["title"])
        set_run_font(run, 15 if section_data["level"] == 1 else 13, bold=True)
        add_bookmark(heading, anchors[section_data["id"]], index)

        for block in section_data["blocks"]:
            block_type = block["type"]
            if block_type == "paragraph":
                add_text_paragraph(doc, block["text"])
            elif block_type == "label":
                add_text_paragraph(doc, block["text"], bold=True)
            elif block_type == "steps":
                add_list(doc, block["items"], "List Number")
            elif block_type == "bullets":
                add_list(doc, block["items"], "List Bullet")
            elif block_type == "table":
                add_table(doc, block["headers"], block["rows"])
            elif block_type == "image":
                add_image(doc, block)

    DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)
    for stale_word in DOWNLOAD_DIR.glob("*.docx"):
        stale_word.unlink()
    output_path = DOWNLOAD_DIR / OUTPUT_NAME
    mark_fields_dirty(doc)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    build_document()

